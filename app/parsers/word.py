# =============================================================================
# 文件作用与架构位置（Word 文档解析器）
# =============================================================================
# 本文件只有 WordParser.parse()。它按 Word XML 中的真实出现顺序遍历段落和表格，解决
# python-docx 分别读取 paragraphs/tables 时可能丢失二者交错顺序的问题。
#
#   .doc/.docx 文件 -> python-docx
#                         |
#                         v
#                  遍历 document body
#                    |             |
#                  段落           表格
#                    |             |
#                    +------合并---+
#                           |
#                    ContentBlock 列表
#                           |
#             full_text + ParsedPage + ParsedDocument
#
# 注意：虽然工厂把 doc 和 docx 都映射到本解析器，但 python-docx 原生主要支持 docx；
# 旧版二进制 doc 是否可读取取决于上游是否已完成格式转换。
# =============================================================================

from typing import List
from app.parsers.base import BaseParser, ParsedDocument, ParsedPage, ContentBlock


class WordParser(BaseParser):
    def parse(self, file_path: str) -> ParsedDocument:
        # 函数内导入较重的 python-docx 依赖，只有处理 Word 时才加载。
        from docx import Document as DocxDocument
        # qn 把 w:p、w:tbl 等带命名空间缩写转换为完整 XML 标签名。
        from docx.oxml.ns import qn
        # Paragraph 可以把底层 XML 段落节点包装成易读取的 Python 对象。
        from docx.text.paragraph import Paragraph

        # 打开 Word 文档。别名 DocxDocument 避免与项目的 Document ORM 模型混淆。
        docx = DocxDocument(file_path)
        # blocks 按原文顺序保存正文块和完整表格块。
        blocks: List[ContentBlock] = []

        # 按文档顺序遍历 body 元素（段落和表格交错）
        # body 遍历得到 XML 表格节点，而具体表格数据从 docx.tables 按同样顺序取得。
        table_index = 0
        # 连续段落先暂存，遇到表格时合并为一个 text 块。
        current_paragraphs: List[str] = []

        # 直接遍历 body 子节点，才能保留“段落 -> 表格 -> 段落”的真实顺序。
        for child in docx.element.body:
            # 段落
            if child.tag == qn("w:p"):
                para = Paragraph(child, docx)
                text = para.text.strip()
                if text:
                    # 空段落没有可检索内容，不加入结果。
                    current_paragraphs.append(text)

            # 表格
            elif child.tag == qn("w:tbl"):
                # 先落盘此表之前的段落
                if current_paragraphs:
                    # 表格前积累的多个段落合并为一个文本块，然后清空缓冲区。
                    blocks.append(ContentBlock(
                        type="text",
                        content="\n".join(current_paragraphs),
                    ))
                    current_paragraphs = []

                # 提取表格，整表作为一个原子块
                if table_index < len(docx.tables):
                    table = docx.tables[table_index]
                    rows = []
                    for row in table.rows:
                        # 同一行的单元格以 | 分隔，使二维结构在纯文本中仍较清晰。
                        cells = [cell.text.strip() for cell in row.cells]
                        rows.append(" | ".join(cells))
                    if rows:
                        blocks.append(ContentBlock(
                            type="table",
                            content="\n".join(rows),
                        ))
                    table_index += 1

        # 落盘末尾的段落
        if current_paragraphs:
            # 文档最后没有表格触发落盘，因此循环结束后必须处理剩余段落。
            blocks.append(ContentBlock(
                type="text",
                content="\n".join(current_paragraphs),
            ))

        # 构建向后兼容的 full_text
        full_text = "\n\n".join(b.content for b in blocks)
        # 最终统一修复乱码和多余空白；blocks 本身仍保留 text/table 类型。
        full_text = self.clean_text(full_text)

        if full_text or blocks:
            doc = ParsedDocument()
            # Word 不是按页面稳定存储内容，当前实现把整份文档视为第 1 页。
            doc.pages.append(ParsedPage(
                page_num=1,
                content=full_text,
                blocks=blocks,
            ))
            return doc

        # 完全没有正文或表格时返回结构合法的空结果。
        return ParsedDocument()
